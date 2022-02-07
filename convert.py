import os
import shutil
import sys

from docx import Document
from docx.text.paragraph import Paragraph

from addons import get_addons
from str_convert import replace
from song import get_songs, add_authors, get_authors

head = True
VERSION = "4.2.1"

if __name__ == '__main__':
    doc = Document(f"docx/Śpiewnik-{VERSION}.docx")
    pars = doc.paragraphs
    sections = dict()

    current_section = None
    for par in doc.paragraphs:
        par: Paragraph
        if par.style.name == "Heading 1":
            sections[par.text] = list()
            current_section = par.text
        elif current_section is not None:
            sections[current_section].append(par)

    if os.path.exists("songbook-online"):
        shutil.rmtree("songbook-online")
    os.mkdir("songbook-online")

    with open("templates/songsWithoutAuthor", "r", encoding="utf-8") as file:
        songs_without_author = [line.strip() for line in file.readlines()]

    for section in sections:
        os.mkdir(f"songbook-online/{section}")
        if section == "Dodatki":
            addons = get_addons(sections[section])
            for addon in addons:
                with open(f"songbook-online/{section}/{addon}.html", "w", encoding="utf-8") as file:
                    file.write(addons[addon])

        else:
            songs = get_songs(sections[section])
            add_authors(songs, get_authors(f"docx/Śpiewnik-{VERSION}.htm"))

            for song in songs:
                if (song.authors is None or len(song.authors) == 0) and song.title not in songs_without_author:
                    print(f"Brak autora w piosence: {song.title}", file=sys.stderr)
                title = song.title.title()
                filename = replace(''.join(x.lower() for x in "-".join(title.split(" ")) if x.isalpha() or x.isnumeric() or x == "-")).encode("ascii", "ignore").decode()
                if filename.isnumeric():
                    filename = "a" + filename
                with open(f"songbook-online/{section}/{filename}.html", "w", encoding="utf-8") as file:
                    if head:
                        file.write("<head><meta charset=\"UTF-8\">")
                        file.write("<link rel=\"stylesheet\" href=\"../../styles.css\"></head>")
                    file.write("<div id=\"song\">")
                    file.write(song.to_html())
                    file.write("</div>")
